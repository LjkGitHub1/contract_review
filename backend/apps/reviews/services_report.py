"""
审核报告生成服务模块
"""
import os
import logging
from typing import Dict, Optional
from django.conf import settings
from django.utils import timezone
from pathlib import Path
from apps.reviews.models import ReviewResult, ReviewOpinion
from apps.contracts.models import Contract

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning('python-docx未安装，Word报告生成功能将不可用')

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning('pymupdf未安装，PDF报告生成功能将不可用')

logger = logging.getLogger(__name__)


class ReportGeneratorService:
    """审核报告生成服务类"""
    
    def __init__(self):
        self.report_dir = Path(settings.MEDIA_ROOT) / 'reports'
        self.report_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_word_report(
        self,
        review_result: ReviewResult,
        contract: Contract
    ) -> str:
        """
        生成Word格式审核报告
        
        Args:
            review_result: 审核结果对象
            contract: 合同对象
            
        Returns:
            str: 报告文件路径
        """
        if not DOCX_AVAILABLE:
            raise Exception('python-docx未安装，无法生成Word报告')
        
        try:
            # 创建文档
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('合同审核报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加报告信息
            doc.add_paragraph(f'报告生成时间：{timezone.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
            doc.add_paragraph(f'合同编号：{contract.contract_no}')
            doc.add_paragraph(f'合同标题：{contract.title}')
            doc.add_paragraph(f'合同类型：{contract.get_contract_type_display()}')
            doc.add_paragraph(f'所属行业：{contract.industry or "通用"}')
            doc.add_paragraph('')
            
            # 添加审核结果概览
            doc.add_heading('一、审核结果概览', 1)
            doc.add_paragraph(f'总体评分：{review_result.overall_score}分')
            doc.add_paragraph(f'风险等级：{review_result.get_risk_level_display()}')
            doc.add_paragraph(f'风险数量：{review_result.risk_count}个')
            doc.add_paragraph('')
            
            # 添加审核摘要
            if review_result.summary:
                doc.add_heading('二、审核摘要', 1)
                doc.add_paragraph(review_result.summary)
                doc.add_paragraph('')
            
            # 添加审核意见详情
            opinions = ReviewOpinion.objects.filter(review_result=review_result).order_by('-risk_level', 'id')
            if opinions.exists():
                doc.add_heading('三、审核意见详情', 1)
                
                for idx, opinion in enumerate(opinions, 1):
                    # 意见标题
                    risk_level_display = opinion.get_risk_level_display()
                    risk_color = {
                        'high': RGBColor(255, 0, 0),  # 红色
                        'medium': RGBColor(255, 165, 0),  # 橙色
                        'low': RGBColor(0, 128, 0)  # 绿色
                    }.get(opinion.risk_level, RGBColor(0, 0, 0))
                    
                    heading = doc.add_heading(f'{idx}. {opinion.get_opinion_type_display()}', 2)
                    for run in heading.runs:
                        run.font.color.rgb = risk_color
                    
                    # 条款信息
                    doc.add_paragraph(f'条款ID：{opinion.clause_id or "未指定"}')
                    if opinion.clause_content:
                        doc.add_paragraph(f'条款内容：{opinion.clause_content[:200]}...')
                    
                    # 意见内容
                    doc.add_paragraph(f'审核意见：{opinion.opinion_content}')
                    
                    # 风险等级
                    risk_para = doc.add_paragraph(f'风险等级：{risk_level_display}')
                    for run in risk_para.runs:
                        if '风险等级' in run.text:
                            run.font.color.rgb = risk_color
                    
                    # 法律依据
                    if opinion.legal_basis:
                        doc.add_paragraph(f'法律依据：{opinion.legal_basis}')
                    
                    # 建议
                    if opinion.suggestion:
                        doc.add_paragraph(f'修改建议：{opinion.suggestion}')
                    
                    doc.add_paragraph('')
            
            # 添加总结
            doc.add_heading('四、总结', 1)
            conclusion = '通过' if review_result.overall_score >= 80 else '需要修改'
            doc.add_paragraph(f'审核结论：{conclusion}')
            doc.add_paragraph('')
            doc.add_paragraph('本报告由AI智能合同审核系统自动生成。')
            
            # 保存文件
            filename = f'report_{review_result.id}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
            filepath = self.report_dir / filename
            doc.save(str(filepath))
            
            # 返回相对路径
            relative_path = f'reports/{filename}'
            return relative_path
            
        except Exception as e:
            logger.error(f'生成Word报告失败: {str(e)}')
            raise
    
    def generate_pdf_report(
        self,
        review_result: ReviewResult,
        contract: Contract
    ) -> str:
        """
        生成PDF格式审核报告
        
        Args:
            review_result: 审核结果对象
            contract: 合同对象
            
        Returns:
            str: 报告文件路径
        """
        if not PDF_AVAILABLE:
            raise Exception('pymupdf未安装，无法生成PDF报告')
        
        try:
            # 创建PDF文档
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)  # A4尺寸
            
            # 设置字体
            font_size = 12
            title_font_size = 20
            
            # 添加标题
            title_rect = fitz.Rect(50, 50, 545, 100)
            page.insert_text(
                title_rect.tl,
                '合同审核报告',
                fontsize=title_font_size,
                color=(0, 0, 0),
                align=1  # 居中
            )
            
            y_pos = 120
            
            # 添加报告信息
            info_lines = [
                f'报告生成时间：{timezone.now().strftime("%Y年%m月%d日 %H:%M:%S")}',
                f'合同编号：{contract.contract_no}',
                f'合同标题：{contract.title}',
                f'合同类型：{contract.get_contract_type_display()}',
                f'所属行业：{contract.industry or "通用"}',
                ''
            ]
            
            for line in info_lines:
                page.insert_text(
                    (50, y_pos),
                    line,
                    fontsize=font_size,
                    color=(0, 0, 0)
                )
                y_pos += 20
            
            # 添加审核结果概览
            y_pos += 10
            page.insert_text(
                (50, y_pos),
                '一、审核结果概览',
                fontsize=font_size + 2,
                color=(0, 0, 0)
            )
            y_pos += 25
            
            overview_lines = [
                f'总体评分：{review_result.overall_score}分',
                f'风险等级：{review_result.get_risk_level_display()}',
                f'风险数量：{review_result.risk_count}个',
                ''
            ]
            
            for line in overview_lines:
                page.insert_text(
                    (50, y_pos),
                    line,
                    fontsize=font_size,
                    color=(0, 0, 0)
                )
                y_pos += 20
            
            # 添加审核摘要
            if review_result.summary:
                y_pos += 10
                page.insert_text(
                    (50, y_pos),
                    '二、审核摘要',
                    fontsize=font_size + 2,
                    color=(0, 0, 0)
                )
                y_pos += 25
                
                # 处理长文本，自动换行
                summary_text = review_result.summary
                words = summary_text.split()
                line = ''
                for word in words:
                    if len(line + word) < 80:
                        line += word + ' '
                    else:
                        page.insert_text(
                            (50, y_pos),
                            line.strip(),
                            fontsize=font_size,
                            color=(0, 0, 0)
                        )
                        y_pos += 20
                        line = word + ' '
                        if y_pos > 800:  # 换页
                            page = doc.new_page(width=595, height=842)
                            y_pos = 50
                
                if line:
                    page.insert_text(
                        (50, y_pos),
                        line.strip(),
                        fontsize=font_size,
                        color=(0, 0, 0)
                    )
                    y_pos += 20
                
                y_pos += 10
            
            # 添加审核意见详情
            opinions = ReviewOpinion.objects.filter(review_result=review_result).order_by('-risk_level', 'id')
            if opinions.exists():
                if y_pos > 750:
                    page = doc.new_page(width=595, height=842)
                    y_pos = 50
                
                page.insert_text(
                    (50, y_pos),
                    '三、审核意见详情',
                    fontsize=font_size + 2,
                    color=(0, 0, 0)
                )
                y_pos += 25
                
                for idx, opinion in enumerate(opinions, 1):
                    if y_pos > 750:
                        page = doc.new_page(width=595, height=842)
                        y_pos = 50
                    
                    # 风险颜色
                    risk_colors = {
                        'high': (1, 0, 0),  # 红色
                        'medium': (1, 0.65, 0),  # 橙色
                        'low': (0, 0.5, 0)  # 绿色
                    }
                    risk_color = risk_colors.get(opinion.risk_level, (0, 0, 0))
                    
                    # 意见标题
                    page.insert_text(
                        (50, y_pos),
                        f'{idx}. {opinion.get_opinion_type_display()}',
                        fontsize=font_size + 1,
                        color=risk_color
                    )
                    y_pos += 20
                    
                    # 意见内容
                    opinion_lines = [
                        f'条款ID：{opinion.clause_id or "未指定"}',
                        f'审核意见：{opinion.opinion_content}',
                        f'风险等级：{opinion.get_risk_level_display()}',
                    ]
                    
                    if opinion.legal_basis:
                        opinion_lines.append(f'法律依据：{opinion.legal_basis}')
                    if opinion.suggestion:
                        opinion_lines.append(f'修改建议：{opinion.suggestion}')
                    
                    for line in opinion_lines:
                        if y_pos > 800:
                            page = doc.new_page(width=595, height=842)
                            y_pos = 50
                        
                        page.insert_text(
                            (50, y_pos),
                            line,
                            fontsize=font_size,
                            color=(0, 0, 0)
                        )
                        y_pos += 20
                    
                    y_pos += 10
            
            # 添加总结
            if y_pos > 750:
                page = doc.new_page(width=595, height=842)
                y_pos = 50
            
            page.insert_text(
                (50, y_pos),
                '四、总结',
                fontsize=font_size + 2,
                color=(0, 0, 0)
            )
            y_pos += 25
            
            conclusion = '通过' if review_result.overall_score >= 80 else '需要修改'
            page.insert_text(
                (50, y_pos),
                f'审核结论：{conclusion}',
                fontsize=font_size,
                color=(0, 0, 0)
            )
            y_pos += 20
            
            page.insert_text(
                (50, y_pos),
                '本报告由AI智能合同审核系统自动生成。',
                fontsize=font_size - 2,
                color=(0.5, 0.5, 0.5)
            )
            
            # 保存文件
            filename = f'report_{review_result.id}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            filepath = self.report_dir / filename
            doc.save(str(filepath))
            doc.close()
            
            # 返回相对路径
            relative_path = f'reports/{filename}'
            return relative_path
            
        except Exception as e:
            logger.error(f'生成PDF报告失败: {str(e)}')
            raise

